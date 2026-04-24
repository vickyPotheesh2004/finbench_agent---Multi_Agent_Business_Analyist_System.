"""
N19 Output Generator — Professional DOCX Report
PDR-BAAAI-001 · Rev 1.0 · Node N19

Purpose:
    Generates professional DOCX report from final answer + all analysis.
    Sections:
        1. Cover page   — company, query, date, confidence
        2. Answer       — final answer with citations
        3. Forensics    — risk score, flags (from N13)
        4. Explainability — SHAP features, causal DAG PNG
        5. Methodology  — PIV rounds, pods used, retrieval info
    CI/CD enforced: zero _rlef_ fields ever appear in output (C9)

Constraints satisfied:
    C1  $0 cost — python-docx is free
    C2  100% local — zero network calls
    C5  seed=42
    C9  _rlef_ fields NEVER in any DOCX output
"""

from __future__ import annotations

import logging
import os
from datetime import datetime
from typing import Dict, List, Optional

logger = logging.getLogger(__name__)

OUTPUT_DIR    = "outputs"
DEFAULT_FNAME = "financial_analysis_report.docx"


class DOCXGenerator:
    """
    N19 Output Generator — creates professional DOCX reports.

    Two usage modes:
        1. gen.generate(state, output_path) → str (path to DOCX)
        2. gen.run(ba_state)               → BAState
    """

    def __init__(self, output_dir: str = OUTPUT_DIR) -> None:
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)

    # ── LangGraph pipeline node ───────────────────────────────────────────────

    def run(self, state) -> object:
        """
        LangGraph N19 node entry point.
        Reads all analysis fields from BAState.
        Writes: state.final_report_path, state.final_answer
        """
        session_id = getattr(state, "session_id", "unknown") or "unknown"
        fname      = f"report_{session_id[:8]}.docx"
        out_path   = os.path.join(self.output_dir, fname)

        # Promote pre-xgb answer to final if not overridden
        final = getattr(state, "xgb_ranked_answer",    "") or ""
        if not final:
            final = getattr(state, "final_answer_pre_xgb", "") or ""

        state.final_answer = final

        path = self.generate(state=state, output_path=out_path)
        state.final_report_path = path

        logger.info("N19 Output: DOCX saved → %s", path)
        return state

    # ── Core generate method ──────────────────────────────────────────────────

    def generate(self, state, output_path: str) -> str:
        """
        Generate DOCX report from BAState.
        Returns path to saved file.
        """
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            logger.warning("python-docx not installed — writing plain text fallback")
            return self._plain_text_fallback(state, output_path)

        doc = Document()
        self._set_margins(doc)

        # ── Section 1: Cover page ─────────────────────────────────────────────
        self._add_cover(doc, state)

        # ── Section 2: Final Answer ───────────────────────────────────────────
        self._add_answer_section(doc, state)

        # ── Section 3: Forensic Risk ──────────────────────────────────────────
        self._add_forensics_section(doc, state)

        # ── Section 4: Explainability ─────────────────────────────────────────
        self._add_explainability_section(doc, state)

        # ── Section 5: Methodology ────────────────────────────────────────────
        self._add_methodology_section(doc, state)

        # C9 gate — never write _rlef_ fields
        self._assert_no_rlef_in_doc(doc)

        os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
        doc.save(output_path)
        return output_path

    # ── Section builders ──────────────────────────────────────────────────────

    def _add_cover(self, doc, state) -> None:
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        company     = getattr(state, "company_name",  "Financial Document") or "Financial Document"
        doc_type    = getattr(state, "doc_type",       "SEC Filing")         or "SEC Filing"
        fiscal_year = getattr(state, "fiscal_year",    "")                   or ""
        confidence  = getattr(state, "confidence_score", 0.0)
        query       = getattr(state, "query",          "")                   or ""
        date_str    = datetime.utcnow().strftime("%B %d, %Y")

        # Title
        title = doc.add_heading("Financial Analysis Report", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.runs[0].font.color.rgb = RGBColor(0x1A, 0x5C, 0x8A)

        # Company + document info
        doc.add_paragraph("")
        info = doc.add_paragraph()
        info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info.add_run(f"{company}  |  {doc_type}").bold = True
        if fiscal_year:
            info.add_run(f"  |  {fiscal_year}")

        # Date + confidence
        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta.add_run(f"Generated: {date_str}  |  Confidence: {confidence:.1%}")

        doc.add_paragraph("")

        # Query
        if query:
            qp = doc.add_paragraph()
            qp.add_run("Query: ").bold = True
            qp.add_run(query[:300])

        doc.add_page_break()

    def _add_answer_section(self, doc, state) -> None:
        final_answer = (
            getattr(state, "final_answer",         "") or
            getattr(state, "final_answer_pre_xgb", "") or
            ""
        )
        citations    = getattr(state, "analyst_citations", []) or []
        low_conf     = getattr(state, "low_confidence",    False)

        doc.add_heading("Analysis Result", 1)

        if low_conf:
            warn = doc.add_paragraph()
            warn.add_run(
                "⚠ LOW CONFIDENCE: This answer required multiple retries and "
                "could not pass all validation checks. Verify against source document."
            ).bold = True

        if final_answer:
            doc.add_paragraph(final_answer)
        else:
            doc.add_paragraph("No answer was generated for this query.")

        if citations:
            doc.add_heading("Citations", 2)
            for cite in citations[:20]:
                doc.add_paragraph(f"• {cite}", style="List Bullet")

    def _add_forensics_section(self, doc, state) -> None:
        risk_score   = getattr(state, "risk_score",       0.0)  or 0.0
        severity     = getattr(state, "anomaly_severity", "low") or "low"
        flags        = getattr(state, "forensic_flags",   [])    or []
        anomaly      = getattr(state, "anomaly_detected", False)

        doc.add_heading("Forensic Risk Analysis", 1)

        # Risk score summary
        risk_para = doc.add_paragraph()
        risk_para.add_run("Risk Score: ").bold = True
        risk_para.add_run(f"{risk_score:.1f} / 100")

        sev_para = doc.add_paragraph()
        sev_para.add_run("Severity: ").bold = True
        sev_para.add_run(severity.upper())

        anom_para = doc.add_paragraph()
        anom_para.add_run("Anomaly Detected: ").bold = True
        anom_para.add_run("YES" if anomaly else "NO")

        if flags:
            doc.add_heading("Forensic Flags", 2)
            for flag in flags[:10]:
                # C9 check — never write _rlef_ content
                if "_rlef_" not in flag:
                    doc.add_paragraph(f"• {flag}", style="List Bullet")
        else:
            doc.add_paragraph("No forensic anomalies detected.")

    def _add_explainability_section(self, doc, state) -> None:
        feature_imp  = getattr(state, "feature_importance", None)
        dag_path     = getattr(state, "causal_dag_path",    None)

        doc.add_heading("Explainability", 1)

        # Feature importance table
        if feature_imp and isinstance(feature_imp, dict):
            doc.add_heading("Top Answer Features (SHAP)", 2)
            table = doc.add_table(rows=1, cols=2)
            table.style = "Table Grid"
            hdr_cells    = table.rows[0].cells
            hdr_cells[0].text = "Feature"
            hdr_cells[1].text = "Importance"
            sorted_fi = sorted(
                feature_imp.items(), key=lambda x: x[1], reverse=True
            )[:10]
            for feat, imp in sorted_fi:
                row_cells         = table.add_row().cells
                row_cells[0].text = feat
                row_cells[1].text = f"{imp:.4f}"
        else:
            doc.add_paragraph("SHAP feature analysis not available.")

        # Causal DAG image
        if dag_path and os.path.exists(dag_path):
            doc.add_heading("Financial Causal DAG", 2)
            try:
                from docx.shared import Inches
                doc.add_picture(dag_path, width=Inches(5.5))
                doc.add_paragraph(
                    "Figure: Revenue → Gross Profit → Operating Income → "
                    "Net Income → EPS causal chain."
                )
            except Exception as e:
                doc.add_paragraph(f"[DAG image unavailable: {e}]")
        else:
            doc.add_paragraph("Causal DAG chart not available.")

    def _add_methodology_section(self, doc, state) -> None:
        piv_round    = getattr(state, "piv_round",          0)     or 0
        win_pod      = getattr(state, "winning_pod",        "")    or ""
        query_type   = getattr(state, "query_type",         "text")
        difficulty   = getattr(state, "query_difficulty",   "medium")
        sniper_hit   = getattr(state, "sniper_hit",         False)
        chunks_count = len(getattr(state, "retrieval_stage_2", []) or [])
        model_ver    = getattr(state, "model_version",
                               "financebench-expert-v1") or "financebench-expert-v1"

        if hasattr(query_type, "value"):
            query_type = query_type.value
        if hasattr(difficulty, "value"):
            difficulty = difficulty.value

        doc.add_heading("Methodology", 1)

        rows = [
            ("Model",              model_ver),
            ("Query Type",         query_type),
            ("Difficulty",         difficulty),
            ("PIV Retries",        str(piv_round)),
            ("Winning Pod",        win_pod or "N/A"),
            ("SniperRAG Hit",      "Yes" if sniper_hit else "No"),
            ("Chunks Retrieved",   str(chunks_count)),
            ("Report Generated",   datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC")),
        ]

        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "Parameter"
        hdr[1].text = "Value"

        for param, value in rows:
            row_cells         = table.add_row().cells
            row_cells[0].text = param
            row_cells[1].text = value

        doc.add_paragraph("")
        footer = doc.add_paragraph()
        footer.add_run(
            "Generated by FinBench Multi-Agent Business Analyst AI · "
            "PDR-BAAAI-001 · $0 cost · 100% local inference"
        ).italic = True

    # ── Helpers ───────────────────────────────────────────────────────────────

    @staticmethod
    def _set_margins(doc) -> None:
        from docx.shared import Inches
        for section in doc.sections:
            section.top_margin    = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin   = Inches(1.2)
            section.right_margin  = Inches(1.2)

    @staticmethod
    def _assert_no_rlef_in_doc(doc) -> None:
        """C9: Scan all paragraphs and table cells for _rlef_ content."""
        for para in doc.paragraphs:
            assert "_rlef_" not in para.text, \
                f"C9 violation: _rlef_ found in paragraph: {para.text[:50]}"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    assert "_rlef_" not in cell.text, \
                        f"C9 violation: _rlef_ found in table cell"

    def _plain_text_fallback(self, state, output_path: str) -> str:
        """Fallback when python-docx is not installed — write .txt."""
        txt_path = output_path.replace(".docx", ".txt")
        final    = (
            getattr(state, "final_answer", "") or
            getattr(state, "final_answer_pre_xgb", "") or
            "No answer generated."
        )
        with open(txt_path, "w", encoding="utf-8") as f:
            f.write(f"Financial Analysis Report\n")
            f.write(f"{'=' * 40}\n")
            f.write(f"Query:  {getattr(state, 'query', '')}\n")
            f.write(f"Answer: {final}\n")
            f.write(f"Risk:   {getattr(state, 'risk_score', 0.0):.1f}/100\n")
        return txt_path


# ── Convenience wrapper ───────────────────────────────────────────────────────

def run_output_generator(state, output_dir: str = OUTPUT_DIR) -> object:
    """Convenience wrapper for LangGraph N19 node."""
    return DOCXGenerator(output_dir=output_dir).run(state)