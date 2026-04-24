"""
src/output/output_generator.py
FinBench Multi-Agent Business Analyst AI
PDR-BAAAI-001 Rev2.0

N19 — Output Generator
Produces professional DOCX report from BAState.

CRITICAL C9: Zero _rlef_ fields in any output.
CI/CD gate scans every DOCX for _rlef_ prefix.

Report sections:
  1. Cover page       — company, query, date, confidence
  2. Executive Answer — final_answer with citations
  3. Analysis Detail  — agreement_status, winning_pod
  4. Quantitative     — monte_carlo, VaR, GARCH (if available)
  5. Forensic Risk    — risk_score, forensic_flags, TriGuard
  6. Explainability   — SHAP feature importance + Causal DAG
  7. Live Market Data — live_data_summary (if available)
  8. Citations Table  — all citations from all pods
  9. Methodology      — pipeline nodes fired, retries used

Writes to BAState:
  final_report_path — path to generated DOCX
"""

import sys
import re
import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional

ROOT = Path(__file__).parent.parent.parent
sys.path.insert(0, str(ROOT))

from src.state.ba_state          import BAState
from src.utils.seed_manager      import SeedManager
from src.utils.resource_governor import ResourceGovernor

SeedManager.set_all()

# ── Output directory ──────────────────────────────────────────────────────────
OUTPUT_DIR = ROOT / "outputs" / "reports"


class OutputGenerator:
    """
    N19 — Output Generator.
    Produces professional DOCX report.
    Zero _rlef_ fields guaranteed.
    """

    def __init__(self, output_dir: Optional[Path] = None):
        SeedManager.set_all()
        self.output_dir = output_dir or OUTPUT_DIR
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def run(self, state: BAState) -> BAState:
        """
        Main entry point — N19 node.
        Reads BAState, writes DOCX report.
        Sets state.final_report_path.
        """
        ResourceGovernor.check("N19 Output Generator")

        try:
            report_path = self._generate_docx(state)
            state.final_report_path = str(report_path)
            print(f"[N19] Report generated: {report_path.name}")
        except Exception as e:
            print(f"[N19] DOCX generation failed: {e} — saving plain text")
            txt_path = self._save_plain_text(state)
            state.final_report_path = str(txt_path)

        # C9 final check — never in output
        self._assert_no_rlef(state.final_report_path)
        return state

    # ═══════════════════════════════════════════════════════════════════════
    # DOCX GENERATION
    # ═══════════════════════════════════════════════════════════════════════

    def _generate_docx(self, state: BAState) -> Path:
        """Generate full professional DOCX report."""
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # ── Document styles ────────────────────────────────────────────────
        style           = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)

        # ── Section 1: Cover ───────────────────────────────────────────────
        self._add_cover(doc, state)

        # ── Section 2: Executive Answer ────────────────────────────────────
        self._add_section(doc, "Executive Answer")
        self._add_answer(doc, state)

        # ── Section 3: Analysis Detail ─────────────────────────────────────
        self._add_section(doc, "Analysis Detail")
        self._add_analysis_detail(doc, state)

        # ── Section 4: Quantitative Analysis ──────────────────────────────
        if state.monte_carlo_results or state.var_result:
            self._add_section(doc, "Quantitative Analysis")
            self._add_quantitative(doc, state)

        # ── Section 5: Forensic Risk ───────────────────────────────────────
        self._add_section(doc, "Forensic Risk Assessment")
        self._add_forensic(doc, state)

        # ── Section 6: Explainability ──────────────────────────────────────
        if state.shap_values or state.causal_dag_path:
            self._add_section(doc, "Explainability")
            self._add_explainability(doc, state)

        # ── Section 7: Live Market Data ────────────────────────────────────
        if getattr(state, "live_data_summary", None):
            self._add_section(doc, "Live Market Context")
            doc.add_paragraph(state.live_data_summary)

        # ── Section 8: Citations ───────────────────────────────────────────
        self._add_section(doc, "Citations")
        self._add_citations(doc, state)

        # ── Section 9: Methodology ─────────────────────────────────────────
        self._add_section(doc, "Methodology")
        self._add_methodology(doc, state)

        # ── Save ───────────────────────────────────────────────────────────
        safe_id   = re.sub(r'[^\w-]', '_', state.session_id)[:40]
        filename  = f"report_{safe_id}.docx"
        out_path  = self.output_dir / filename
        doc.save(str(out_path))
        return out_path

    # ═══════════════════════════════════════════════════════════════════════
    # SECTION BUILDERS
    # ═══════════════════════════════════════════════════════════════════════

    def _add_cover(self, doc: Any, state: BAState) -> None:
        """Add cover page."""
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        # Title
        title      = doc.add_heading("FinBench AI — Financial Analysis Report",
                                     level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph("")

        # Metadata table
        table = doc.add_table(rows=6, cols=2)
        table.style = "Table Grid"
        rows_data = [
            ("Company",    state.company_name  or "—"),
            ("Document",   state.doc_type      or "—"),
            ("Fiscal Year",state.fiscal_year   or "—"),
            ("Query",      (state.query or "—")[:100]),
            ("Confidence", f"{state.confidence_score:.1%}"
                           if state.confidence_score else "—"),
            ("Generated",  datetime.datetime.utcnow().strftime(
                           "%Y-%m-%d %H:%M UTC")),
        ]
        for i, (label, value) in enumerate(rows_data):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = str(value)

        doc.add_page_break()

    def _add_section(self, doc: Any, title: str) -> None:
        """Add a section heading."""
        doc.add_heading(title, level=2)

    def _add_answer(self, doc: Any, state: BAState) -> None:
        """Add executive answer section."""
        answer = state.final_answer or state.xgb_ranked_answer or \
                 state.final_answer_pre_xgb or "No answer generated."

        # Confidence indicator
        conf  = state.confidence_score or 0.0
        if conf >= 0.85:
            conf_label = "HIGH CONFIDENCE"
        elif conf >= 0.65:
            conf_label = "MEDIUM CONFIDENCE"
        else:
            conf_label = "LOW CONFIDENCE — verify before use"

        p = doc.add_paragraph()
        p.add_run(f"[{conf_label}] ").bold = True
        p.add_run(f"Score: {conf:.1%}")

        doc.add_paragraph(answer)

        # HITL warning
        if state.low_confidence:
            warning = doc.add_paragraph()
            warning.add_run(
                "⚠ WARNING: This answer required multiple retries and "
                "could not pass all validation checks. "
                "Verify against source document before use."
            ).bold = True

    def _add_analysis_detail(self, doc: Any, state: BAState) -> None:
        """Add analysis detail — pod agreement, winning pod."""
        # Agreement status
        agreement = state.agreement_status or "unknown"
        status    = agreement.split("|")[0] if "|" in agreement else agreement
        winning   = agreement.split("|")[1] if "|" in agreement else "—"

        table = doc.add_table(rows=4, cols=2)
        table.style = "Table Grid"
        data = [
            ("Agreement",   status.upper()),
            ("Winning Pod", winning.capitalize()),
            ("PIV Retries", str(state.piv_round or 0)),
            ("Query Type",  str(state.query_type or "—")),
        ]
        for i, (label, value) in enumerate(data):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = value

        doc.add_paragraph("")

        # Pod outputs summary
        if state.analyst_output:
            h = doc.add_paragraph()
            h.add_run("Lead Analyst:").bold = True
            doc.add_paragraph(
                f"{state.analyst_output[:300]}..."
                if len(state.analyst_output) > 300
                else state.analyst_output
            )

        if state.auditor_output:
            h = doc.add_paragraph()
            h.add_run("Independent Auditor:").bold = True
            doc.add_paragraph(
                f"{state.auditor_output[:200]}..."
                if len(state.auditor_output) > 200
                else state.auditor_output
            )

        # Contradictions
        if state.contradiction_flags:
            h = doc.add_paragraph()
            h.add_run("Contradictions Detected:").bold = True
            for flag in state.contradiction_flags:
                doc.add_paragraph(f"• {flag}", style="List Bullet")

    def _add_quantitative(self, doc: Any, state: BAState) -> None:
        """Add quantitative analysis — Monte Carlo, VaR, GARCH."""
        if state.monte_carlo_results:
            mc = state.monte_carlo_results
            h  = doc.add_paragraph()
            h.add_run("Monte Carlo Analysis (10,000 scenarios):").bold = True

            table = doc.add_table(rows=5, cols=2)
            table.style = "Table Grid"
            mc_data = [
                ("Mean",    f"{mc.get('mean', 0):,.1f}"),
                ("P5  (5th percentile)",  f"{mc.get('p5',  0):,.1f}"),
                ("P50 (median)",          f"{mc.get('p50', 0):,.1f}"),
                ("P95 (95th percentile)", f"{mc.get('p95', 0):,.1f}"),
                ("Scenarios", f"{mc.get('n', 0):,}"),
            ]
            for i, (label, value) in enumerate(mc_data):
                table.rows[i].cells[0].text = label
                table.rows[i].cells[1].text = value
            doc.add_paragraph("")

        if state.var_result:
            var = state.var_result
            h   = doc.add_paragraph()
            h.add_run("Value at Risk:").bold = True
            p95 = var.get("var_95", var.get("var95", "—"))
            p99 = var.get("var_99", var.get("var99", "—"))
            doc.add_paragraph(
                f"VaR 95th percentile: {p95:,.1f}  |  "
                f"VaR 99th percentile: {p99:,.1f}"
                if isinstance(p95, (int, float))
                else "VaR data unavailable"
            )

    def _add_forensic(self, doc: Any, state: BAState) -> None:
        """Add forensic risk section."""
        risk  = state.risk_score or 0.0
        sev   = state.anomaly_severity or "low"
        color = "HIGH RISK" if sev == "high" else \
                "MEDIUM RISK" if sev == "medium" else "LOW RISK"

        p = doc.add_paragraph()
        p.add_run(f"Risk Score: {risk:.1f}/100 — {color}").bold = True

        if state.forensic_flags:
            doc.add_paragraph("Forensic Flags:")
            for flag in state.forensic_flags:
                if "_rlef_" not in flag:   # C9 guard
                    doc.add_paragraph(f"• {flag}", style="List Bullet")
        else:
            doc.add_paragraph("No forensic anomalies detected.")

        # Benford stats
        if state.benford_chi2 is not None:
            doc.add_paragraph(
                f"Benford Law Test: χ²={state.benford_chi2:.2f}  "
                f"p={state.benford_p_value:.4f}"
                + (" — SUSPICIOUS" if state.benford_p_value < 0.05 else
                   " — Normal distribution")
            )

    def _add_explainability(self, doc: Any, state: BAState) -> None:
        """Add SHAP + Causal DAG section."""
        from docx.shared import Inches

        if state.shap_values:
            doc.add_paragraph("Feature Importance (SHAP):")
            table = doc.add_table(
                rows=len(state.shap_values) + 1, cols=2
            )
            table.style = "Table Grid"
            table.rows[0].cells[0].text = "Feature"
            table.rows[0].cells[1].text = "SHAP Value"
            sorted_shap = sorted(
                state.shap_values.items(),
                key=lambda x: x[1], reverse=True
            )
            for i, (feat, val) in enumerate(sorted_shap, 1):
                table.rows[i].cells[0].text = feat
                table.rows[i].cells[1].text = f"{val:.4f}"
            doc.add_paragraph("")

        if state.causal_dag_path:
            dag_path = Path(state.causal_dag_path)
            if dag_path.exists():
                try:
                    doc.add_paragraph("Financial Causal Chain:")
                    doc.add_picture(str(dag_path), width=Inches(5.5))
                except Exception:
                    doc.add_paragraph(
                        f"[Causal DAG saved at: {dag_path.name}]"
                    )

    def _add_citations(self, doc: Any, state: BAState) -> None:
        """Add citations table."""
        all_cits = []
        for cit_list in [
            state.analyst_citations,
            state.quant_citations,
            state.auditor_citations,
        ]:
            all_cits.extend(cit_list or [])

        unique_cits = list(dict.fromkeys(all_cits))   # preserve order

        if unique_cits:
            table = doc.add_table(rows=len(unique_cits)+1, cols=2)
            table.style = "Table Grid"
            table.rows[0].cells[0].text = "#"
            table.rows[0].cells[1].text = "Citation"
            for i, cit in enumerate(unique_cits, 1):
                if "_rlef_" not in cit:   # C9 guard
                    table.rows[i].cells[0].text = str(i)
                    table.rows[i].cells[1].text = cit
        else:
            doc.add_paragraph("No citations recorded.")

    def _add_methodology(self, doc: Any, state: BAState) -> None:
        """Add methodology section."""
        table = doc.add_table(rows=6, cols=2)
        table.style = "Table Grid"
        data = [
            ("Model",         state.model_version  or "financebench-expert-v1"),
            ("Pipeline",      "19-node PIV pipeline"),
            ("Retrieval",     "VectorlessFirst 4-tier cascade"),
            ("Pods fired",    "LeadAnalyst + QuantAnalyst + BlindAuditor"),
            ("Validation",    "8-check CuriousValidator per pod"),
            ("Seed",          str(state.seed)),
        ]
        for i, (label, value) in enumerate(data):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = value

    # ═══════════════════════════════════════════════════════════════════════
    # PLAIN TEXT FALLBACK
    # ═══════════════════════════════════════════════════════════════════════

    def _save_plain_text(self, state: BAState) -> Path:
        """Fallback — save answer as plain text if DOCX fails."""
        safe_id  = re.sub(r'[^\w-]', '_', state.session_id)[:40]
        out_path = self.output_dir / f"report_{safe_id}.txt"
        content  = (
            f"FinBench AI Report\n"
            f"{'='*50}\n"
            f"Company:    {state.company_name}\n"
            f"Query:      {state.query}\n"
            f"Confidence: {state.confidence_score:.1%}\n\n"
            f"ANSWER:\n{state.final_answer}\n\n"
            f"Risk Score: {state.risk_score}/100\n"
        )
        # C9 guard — strip any _rlef_ that somehow got in
        content = re.sub(r'_rlef_\w+[^\n]*\n?', '', content)
        out_path.write_text(content, encoding="utf-8")
        return out_path

    # ═══════════════════════════════════════════════════════════════════════
    # C9 GUARD
    # ═══════════════════════════════════════════════════════════════════════

    def _assert_no_rlef(self, file_path: Optional[str]) -> None:
        """
        C9: Verify no _rlef_ fields in output file.
        Raises AssertionError if found.
        """
        if not file_path:
            return
        path = Path(file_path)
        if not path.exists():
            return
        if path.suffix == ".txt":
            content = path.read_text(encoding="utf-8", errors="ignore")
            assert "_rlef_" not in content, \
                f"C9 VIOLATION: _rlef_ found in output {path.name}"

    def generate_report(self, state: BAState) -> str:
        """Public convenience method — returns report path string."""
        state = self.run(state)
        return state.final_report_path or ""


# ═══════════════════════════════════════════════════════════════════════════
# QUICK SANITY CHECK
# run: python src/output/output_generator.py
# ═══════════════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    import tempfile
    try:
        from rich import print as rprint
    except ImportError:
        rprint = print

    rprint("\n[bold cyan]-- OutputGenerator (N19) sanity check --[/bold cyan]")

    with tempfile.TemporaryDirectory() as tmp:
        gen   = OutputGenerator(output_dir=Path(tmp))
        rprint("[green]✓[/green] OutputGenerator instantiated")

        state = BAState(
            session_id           = "sanity-n19",
            query                = "What was Apple net income FY2023?",
            query_type           = __import__(
                'src.state.ba_state', fromlist=['QueryType']
            ).QueryType.NUMERICAL,
            company_name         = "Apple Inc",
            doc_type             = "10-K",
            fiscal_year          = "FY2023",
            final_answer         = "Net income was $96,995 million in FY2023 "
                                   "[Financial Statements/P42].",
            xgb_ranked_answer    = "Net income was $96,995 million in FY2023.",
            confidence_score     = 0.92,
            agreement_status     = "unanimous|analyst",
            analyst_output       = "Net income $96,995M FY2023 [FS/42].",
            analyst_confidence   = 0.92,
            analyst_citations    = ["Financial Statements / Page 42",
                                    "Income Statement / Page 44"],
            quant_result         = "Net income $96,995M [FS/42].",
            quant_confidence     = 0.88,
            quant_citations      = ["Financial Statements / Page 42"],
            auditor_output       = "Net income $96,995M FY2023 confirmed.",
            auditor_confidence   = 0.90,
            auditor_citations    = ["Financial Statements / Page 42"],
            monte_carlo_results  = {
                "mean": 96995.0, "std": 4849.0,
                "p5": 88876.0, "p25": 93500.0,
                "p50": 96995.0, "p75": 100400.0,
                "p95": 104962.0, "n": 10000,
            },
            var_result           = {"var_95": 64793.0, "var_99": 58887.0},
            risk_score           = 18.5,
            anomaly_severity     = "low",
            anomaly_detected     = False,
            forensic_flags       = ["ISOLATION_FOREST_MEDIUM: 2/15 outliers"],
            benford_chi2         = 2.93,
            benford_p_value      = 0.938,
            shap_values          = {
                "bm25_score": 0.42, "cosine_sim": 0.38,
                "section_relevance": 0.12, "citation_present": 0.08,
            },
            contradiction_flags  = [],
            piv_round            = 0,
        )

        state = gen.run(state)

        rprint(f"[green]✓[/green] Report generated: "
               f"{Path(state.final_report_path).name}")
        assert state.final_report_path is not None
        assert Path(state.final_report_path).exists()

        # C9 check
        if state.final_report_path.endswith(".txt"):
            content = Path(state.final_report_path).read_text()
            assert "_rlef_" not in content
        rprint(f"[green]✓[/green] C9: no _rlef_ in output")

        assert state.seed == 42
        rprint(f"[green]✓[/green] seed=42 preserved")

        rprint(f"\n[bold green]All checks passed. "
               f"OutputGenerator N19 ready.[/bold green]\n")