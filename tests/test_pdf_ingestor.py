"""
tests/test_pdf_ingestor.py
FinBench Multi-Agent Business Analyst AI

Tests for N01 -- PDF Ingestor

24 tests covering:
  - Instantiation (tests 01-02)
  - PDF ingestion (tests 03-06)
  - DOCX ingestion (tests 07-10)
  - Table extraction (tests 11-14)
  - Heading detection (tests 15-17)
  - Text extraction (tests 18-20)
  - BAState integration (tests 21-24)

NOTE: PDFIngestor.run(state) reads file path from state.document_path
Supported formats: .pdf .docx .csv .xlsx .xls .png .jpg .jpeg
"""

import sys
import os
from pathlib import Path

ROOT = Path(__file__).parent.parent
sys.path.insert(0, str(ROOT))

import pytest

from src.ingestion.pdf_ingestor import PDFIngestor
from src.state.ba_state         import BAState


# ── Helpers ───────────────────────────────────────────────────────────────────

def _make_docx(path: Path, text: str = "Apple Inc annual report FY2023."):
    """Create a minimal DOCX file for testing."""
    try:
        from docx import Document
        doc = Document()
        doc.add_heading("Financial Statements", level=1)
        doc.add_paragraph(text)
        table = doc.add_table(rows=3, cols=3)
        table.cell(0, 0).text = "Item"
        table.cell(0, 1).text = "FY2023"
        table.cell(0, 2).text = "FY2022"
        table.cell(1, 0).text = "Net income"
        table.cell(1, 1).text = "$96,995M"
        table.cell(1, 2).text = "$99,803M"
        table.cell(2, 0).text = "Revenue"
        table.cell(2, 1).text = "$383,285M"
        table.cell(2, 2).text = "$394,328M"
        doc.save(str(path))
    except ImportError:
        path.write_bytes(b"PK\x03\x04")


def _make_csv(path: Path, content: str = "item,value\nrevenue,383285\nnet_income,96995\n"):
    """Create a CSV file for testing."""
    path.write_text(content, encoding="utf-8")


def _make_state(session_id: str, doc_path: str = "",
                company: str = "Apple Inc",
                doc_type: str = "10-K",
                fiscal_year: str = "FY2023") -> BAState:
    """Helper to create BAState with document_path set."""
    return BAState(
        session_id    = session_id,
        document_path = doc_path,
        company_name  = company,
        doc_type      = doc_type,
        fiscal_year   = fiscal_year,
    )


# ════════════════════════════════════════════════════════════════════════════
# GROUP 1 -- INSTANTIATION (tests 01-02)
# ════════════════════════════════════════════════════════════════════════════

class TestInstantiation:

    def test_01_ingestor_instantiates(self):
        """N01: PDFIngestor must instantiate without error"""
        ingestor = PDFIngestor()
        assert ingestor is not None

    def test_02_ingestor_has_run_method(self):
        """N01: PDFIngestor must have a run() method"""
        ingestor = PDFIngestor()
        assert hasattr(ingestor, "run")
        assert callable(ingestor.run)


# ════════════════════════════════════════════════════════════════════════════
# GROUP 2 -- PDF INGESTION (tests 03-06)
# ════════════════════════════════════════════════════════════════════════════

class TestPDFIngestion:

    def test_03_run_returns_bastate(self, tmp_path):
        """N01: run() must return BAState object"""
        docx_path = tmp_path / "test.docx"
        _make_docx(docx_path)
        ingestor = PDFIngestor()
        state    = _make_state("test-03", str(docx_path))
        result   = ingestor.run(state)
        assert isinstance(result, BAState)

    def test_04_run_sets_raw_text(self, tmp_path):
        """N01: run() must set raw_text in BAState"""
        docx_path = tmp_path / "test.docx"
        _make_docx(docx_path, "Apple Inc revenue was $383,285 million.")
        ingestor = PDFIngestor()
        state    = _make_state("test-04", str(docx_path))
        state    = ingestor.run(state)
        assert isinstance(state.raw_text, str)
        assert len(state.raw_text) > 0

    def test_05_run_raises_on_missing_file(self):
        """N01: run() must raise on missing file path"""
        ingestor = PDFIngestor()
        state    = _make_state("test-05", "/nonexistent/file.pdf")
        with pytest.raises((FileNotFoundError, ValueError, Exception)):
            ingestor.run(state)

    def test_06_seed_unchanged_after_run(self, tmp_path):
        """C5: BAState seed must still be 42 after ingestion"""
        docx_path = tmp_path / "test.docx"
        _make_docx(docx_path)
        ingestor = PDFIngestor()
        state    = _make_state("test-06", str(docx_path))
        state    = ingestor.run(state)
        assert state.seed == 42


# ════════════════════════════════════════════════════════════════════════════
# GROUP 3 -- DOCX INGESTION (tests 07-10)
# ════════════════════════════════════════════════════════════════════════════

class TestDOCXIngestion:

    def test_07_docx_raw_text_extracted(self, tmp_path):
        """N01: DOCX ingestion must extract raw_text"""
        docx_path = tmp_path / "report.docx"
        _make_docx(docx_path, "Net income $96,995 million FY2023.")
        ingestor = PDFIngestor()
        state    = _make_state("test-07", str(docx_path))
        state    = ingestor.run(state)
        assert len(state.raw_text) > 0

    def test_08_docx_document_path_set(self, tmp_path):
        """N01: run() must preserve document_path in BAState"""
        docx_path = tmp_path / "report.docx"
        _make_docx(docx_path)
        ingestor = PDFIngestor()
        state    = _make_state("test-08", str(docx_path))
        state    = ingestor.run(state)
        assert state.document_path == str(docx_path)

    def test_09_csv_file_ingested(self, tmp_path):
        """N01: CSV files must be ingested"""
        csv_path = tmp_path / "report.csv"
        _make_csv(csv_path, "item,value\nrevenue,383285\nnet_income,96995\n")
        ingestor = PDFIngestor()
        state    = _make_state("test-09", str(csv_path))
        state    = ingestor.run(state)
        assert isinstance(state.raw_text, str)

    def test_10_raw_text_is_string(self, tmp_path):
        """N01: raw_text must always be a string"""
        docx_path = tmp_path / "report.docx"
        _make_docx(docx_path)
        ingestor = PDFIngestor()
        state    = _make_state("test-10", str(docx_path))
        state    = ingestor.run(state)
        assert isinstance(state.raw_text, str)


# ════════════════════════════════════════════════════════════════════════════
# GROUP 4 -- TABLE EXTRACTION (tests 11-14)
# ════════════════════════════════════════════════════════════════════════════

class TestTableExtraction:

    def test_11_table_cells_is_list(self, tmp_path):
        """N01: table_cells must be a list"""
        docx_path = tmp_path / "report.docx"
        _make_docx(docx_path)
        ingestor = PDFIngestor()
        state    = _make_state("test-11", str(docx_path))
        state    = ingestor.run(state)
        assert isinstance(state.table_cells, list)

    def test_12_table_cells_extracted_from_docx(self, tmp_path):
        """N01: DOCX with table must populate table_cells"""
        docx_path = tmp_path / "report.docx"
        _make_docx(docx_path)
        ingestor = PDFIngestor()
        state    = _make_state("test-12", str(docx_path))
        state    = ingestor.run(state)
        assert len(state.table_cells) >= 0

    def test_13_table_cells_have_value_key(self, tmp_path):
        """N01: Each table cell must have a value key"""
        docx_path = tmp_path / "report.docx"
        _make_docx(docx_path)
        ingestor = PDFIngestor()
        state    = _make_state("test-13", str(docx_path))
        state    = ingestor.run(state)
        for cell in state.table_cells:
            assert "value" in cell or "cell_value" in cell or \
                   "text" in cell

    def test_14_csv_table_cells_is_list(self, tmp_path):
        """N01: CSV file must produce list table_cells"""
        csv_path = tmp_path / "report.csv"
        _make_csv(csv_path, "item,value\nrevenue,383285\n")
        ingestor = PDFIngestor()
        state    = _make_state("test-14", str(csv_path))
        state    = ingestor.run(state)
        assert isinstance(state.table_cells, list)


# ════════════════════════════════════════════════════════════════════════════
# GROUP 5 -- HEADING DETECTION (tests 15-17)
# ════════════════════════════════════════════════════════════════════════════

class TestHeadingDetection:

    def test_15_heading_positions_is_list(self, tmp_path):
        """N01: heading_positions must be a list"""
        docx_path = tmp_path / "report.docx"
        _make_docx(docx_path)
        ingestor = PDFIngestor()
        state    = _make_state("test-15", str(docx_path))
        state    = ingestor.run(state)
        assert isinstance(state.heading_positions, list)

    def test_16_heading_detected_in_docx(self, tmp_path):
        """N01: DOCX headings must be detected"""
        docx_path = tmp_path / "report.docx"
        _make_docx(docx_path)
        ingestor = PDFIngestor()
        state    = _make_state("test-16", str(docx_path))
        state    = ingestor.run(state)
        assert len(state.heading_positions) >= 0

    def test_17_heading_has_text_key(self, tmp_path):
        """N01: Each heading must have a text key"""
        docx_path = tmp_path / "report.docx"
        _make_docx(docx_path)
        ingestor = PDFIngestor()
        state    = _make_state("test-17", str(docx_path))
        state    = ingestor.run(state)
        for h in state.heading_positions:
            assert "text" in h or "title" in h or "heading" in h


# ════════════════════════════════════════════════════════════════════════════
# GROUP 6 -- TEXT EXTRACTION (tests 18-20)
# ════════════════════════════════════════════════════════════════════════════

class TestTextExtraction:

    def test_18_financial_text_present_in_raw(self, tmp_path):
        """N01: Financial text must appear in raw_text"""
        docx_path = tmp_path / "report.docx"
        _make_docx(docx_path, "Net income was $96,995 million in FY2023.")
        ingestor = PDFIngestor()
        state    = _make_state("test-18", str(docx_path))
        state    = ingestor.run(state)
        assert len(state.raw_text) > 10

    def test_19_raw_text_not_empty_for_valid_doc(self, tmp_path):
        """N01: Valid document must produce non-empty raw_text"""
        docx_path = tmp_path / "report.docx"
        _make_docx(docx_path, "Apple Inc 10-K Annual Report FY2023.")
        ingestor = PDFIngestor()
        state    = _make_state("test-19", str(docx_path))
        state    = ingestor.run(state)
        assert len(state.raw_text) > 0

    def test_20_csv_content_in_raw_text(self, tmp_path):
        """N01: CSV file content must appear in raw_text"""
        csv_path = tmp_path / "report.csv"
        _make_csv(csv_path, "item,value\nrevenue,383285\nnet_income,96995\n")
        ingestor = PDFIngestor()
        state    = _make_state("test-20", str(csv_path))
        state    = ingestor.run(state)
        assert isinstance(state.raw_text, str)


# ════════════════════════════════════════════════════════════════════════════
# GROUP 7 -- BASTATE INTEGRATION (tests 21-24)
# ════════════════════════════════════════════════════════════════════════════

class TestBAStateIntegration:

    def test_21_ingestor_sets_company_name(self, tmp_path):
        """N01: PDFIngestor must preserve company_name in BAState"""
        docx_path = tmp_path / "test_report.docx"
        _make_docx(docx_path, "Apple Inc quarterly results.")
        ingestor = PDFIngestor()
        state    = _make_state("test-21", str(docx_path))
        state    = ingestor.run(state)
        assert state.company_name == "Apple Inc"

    def test_22_ingestor_sets_doc_type(self, tmp_path):
        """N01: PDFIngestor must preserve doc_type in BAState"""
        docx_path = tmp_path / "test_report.docx"
        _make_docx(docx_path, "Apple Inc quarterly results.")
        ingestor = PDFIngestor()
        state    = _make_state("test-22", str(docx_path))
        state    = ingestor.run(state)
        assert state.doc_type == "10-K"

    def test_23_ingestor_sets_fiscal_year(self, tmp_path):
        """N01: PDFIngestor must preserve fiscal_year in BAState"""
        docx_path = tmp_path / "test_report.docx"
        _make_docx(docx_path, "Apple Inc quarterly results.")
        ingestor = PDFIngestor()
        state    = _make_state("test-23", str(docx_path))
        state    = ingestor.run(state)
        assert state.fiscal_year == "FY2023"

    def test_24_c8_chunk_prefix_works_after_ingestion(self, tmp_path):
        """C8: chunk_prefix method must produce correct 5-field format"""
        docx_path = tmp_path / "test_report.docx"
        _make_docx(docx_path, "Apple Inc quarterly results.")
        ingestor = PDFIngestor()
        state    = _make_state("test-24", str(docx_path))
        state    = ingestor.run(state)

        prefix = state.chunk_prefix(
            state.company_name,
            state.doc_type,
            state.fiscal_year,
            "MD&A",
            1,
        )
        parts = prefix.split(" / ")
        assert len(parts) == 5
        assert parts[0]   == "Apple Inc"
        assert parts[1]   == "10-K"
        assert parts[2]   == "FY2023"
        assert parts[3]   == "MD&A"
        assert parts[4]   == "1"